import os
from omegaconf import OmegaConf
import pandas as pd
import shutil
import datetime as dt
from docx import Document
import re

class PstDB:
    def __init__(self, data_pth, issue, year):
        self.data_pth = data_pth
        year_issue_pth = os.path.join(data_pth, 'yaml/year_issue.yaml')
        issue_article_pth = os.path.join(data_pth, 'yaml/issue_article.yaml')
        self.authors_pth = os.path.join(data_pth, 'yaml/authors.yaml')
        self.affiliations_pth = os.path.join(data_pth, 'yaml/affiliations.yaml')
        self.degrees_pth = os.path.join(data_pth, 'yaml/degrees.yaml')
        self.sections_pth = os.path.join(data_pth, 'yaml/sections.yaml')
        self.edn_pth = os.path.join(data_pth, 'yaml/edn.yaml')
        self.edn = OmegaConf.load(self.edn_pth)
        #################
        self.year_issue = OmegaConf.load(year_issue_pth)
        self.issue_article = OmegaConf.load(issue_article_pth)

        self.issue_article_pth = os.path.join(data_pth, 'yaml/issue_article.yaml')
        self.issue_article = OmegaConf.load(self.issue_article_pth)

        self.j_ru = 'Петербургская социология сегодня'
        self.j_en = 'St. Petersburg Sociology Today'

        self.issue = issue
        self.year = year

        self.art_in_iss = None


    def make_zip(self, dir_list):
        curr_dt = dt.datetime.now().strftime('%Y%m%d_%H%M%S')
        for dir_name in dir_list:
            output_filename = os.path.join(self.data_pth, 'a', f'{dir_name}_{curr_dt}')
            shutil.make_archive(output_filename, 'zip', dir_name)

    def add_article(self, art_fn):
        self.make_zip(dir_list=['articles', 'data/yaml']) #######################
        art = OmegaConf.load(art_fn)

        # authors
        authors = OmegaConf.load(self.authors_pth)
        affiliations = OmegaConf.load(self.affiliations_pth)
        for i in art['authors']:
            if (i['id'] is not None) & (i['name'] is not None):
                pass
            if i['id']:
                a_data = authors[i['id']]
                i['name'] = a_data['name']
                i['patronymic'] = a_data['patronymic']
                i['surname'] = a_data['surname']
                i['orcid'] = a_data['orcid']
            else:
                a_ids = [k for k in authors]
                next_id = max(a_ids) + 1
                i['id'] = next_id
                authors[next_id] = {
                    'name': i['name'],
                    'patronymic': i['patronymic'],
                    'surname': i['surname'],
                    'orcid': i['orcid'],
                    'spin': i['spin'],
                }
            # aff
            print(affiliations)
            for aff in i['affs']:
                if aff['id']:
                    aff_name = affiliations[aff['id']]
                    aff['aff_name'] = aff_name['aff_name']
                else:
                    aff_ids = [k for k in affiliations]
                    next_id = max(aff_ids) + 1
                    aff['id'] = next_id
                    affiliations[next_id] = {
                        'aff_name': aff['aff_name']
                    }
            # position

            # degree
            degrees = OmegaConf.load(self.degrees_pth)
            degree = i['degree']
            for d in degree:
                if d in degrees.keys():
                    i['degree'] = degrees[d]
        # section
        sections = OmegaConf.load(self.sections_pth)
        section_id = art['section']['id']
        art['section']['sect_name'] = sections[section_id]

        with open(self.affiliations_pth, 'w') as fp:
            OmegaConf.save(affiliations, fp.name)
        with open(self.authors_pth, 'w') as fp:
            OmegaConf.save(authors, fp.name)
        with open(f"articles/a{art['id']}.yaml", 'w') as fp:
            OmegaConf.save(art, fp.name)

    def print_author(self):
        authors = OmegaConf.load(self.authors_pth) ############################
        author_data_lst = [[k, v['surname'][0], v['name'][0], v['patronymic'][0]] for k, v in authors.items()]
        authors_df = pd.DataFrame(author_data_lst)
        print(authors_df.sort_values(1))

    def get_articles(self, folder):
        curr_art = []
        for r, d, fs in os.walk(folder):
            for f in fs:
                print(f)
                a = OmegaConf.load(os.path.join(r, f))
                curr_art.append(a)
        return curr_art

    def get_statistics(self, issue):
        curr_art = self.get_articles(folder='articles')
        vol = 0
        num_art = 0
        for i in self.issue_article[issue]:
            for ca in curr_art:
                if i == ca['id']:
                    vol += ca['vol']
                    num_art += 1
        al = round(vol / 40000, 1)
        print(vol, ' - ', al, ' - ', f'{round(al / 0.08, 1)} %')

    def get_art_in_iss(self, issue): #################################
        curr_art = self.get_articles(folder='articles') ###################
        art_in_iss = []
        for i in self.issue_article[issue]:
            for ca in curr_art:
                if i == ca['id']:
                    art_in_iss.append(ca)

        for i in art_in_iss:
            autors = [a['surname'][0] for a in i['authors']][0]
            print(i['id'], autors, i['title'][0])
        self.art_in_iss = art_in_iss

    def add_edn2art(self, issue):
        curr_art = self.get_articles(folder='articles')
        # print(curr_art)
        for n, i in enumerate(self.issue_article[issue]):
            for ca in curr_art:
                if i == ca['id']:
                    ca['edn'] = self.edn[issue][n + 1]
                    print(i, ca['edn'])
                    with open(f"articles/a{ca['id']}.yaml", 'w') as fp:
                        OmegaConf.save(ca, fp.name)

    def add_date2art(self, issue, date_field, date4insert):
        curr_art = self.get_articles(folder='articles')
        print(curr_art)
        # accepted
        for n, i in enumerate(self.issue_article[issue]):
            for ca in curr_art:
                if i == ca['id']:
                    if ca[date_field] is None:
                        ca[date_field] = date4insert
                    with open(f"articles/a{ca['id']}.yaml", 'w') as fp:
                        OmegaConf.save(ca, fp.name)

    def empty_spaces(self, issue):
        curr_art = self.get_articles(folder='articles')
        fields4checking = ['udc', 'received', 'revised', 'received']
        for n, i in enumerate(self.issue_article[issue]):
            for ca in curr_art:
                if i == ca['id']:
                    print("-" * 10)
                    print("Пропуски в", i)
                    for fc in fields4checking:
                        if ca[fc] is None:
                            print(fc)


    def get_cit(self, rec, fio, lang='ru'):
        cit = ''
        if lang == 'ru':
            cit = f"{fio} {rec['title'][0]} // {self.j_ru}. — {self.year}. — № {self.issue}. — С. {rec['pp']}. — "
        elif lang == 'en':
            cit = f"{fio} {rec['title'][1]}. {self.j_en}. {self.year}. No {self.issue}. P. {rec['pp']}. "
        cit += f"DOI: {rec['doi']}; EDN: {rec['edn']}"
        return cit

    def add_art(self, document, rec, section):
        if section:
            # document.add_heading(section.upper(), level=2)
            document.add_paragraph(section.upper(), style='section')
        document.add_paragraph(f"DOI: {rec['doi']}", style='Аннотация')
        document.add_paragraph(f"EDN: {rec['edn']}", style='Аннотация')
        document.add_paragraph(f"УДК {rec['udc']}", style='Аннотация')

        # authors, authors_num = get_full_name(rec, lang='ru')
        authors, authors_num, affs = authors_affiliation(rec, lang='ru')
        # authors, affs = authors_affiliation(rec)
        # affs = get_affs(rec, lang='ru')
        about = get_about(rec, lang='ru')

        document.add_paragraph(f"{authors}", style='art_author')
        document.add_paragraph(f"{affs}\n{about}", style='art_aff')
        # document.add_paragraph(f"{about}", style='art_title')
        document.add_paragraph(rec['title'][0], style='art_title')

        p = document.add_paragraph('', style='Аннотация')
        p.add_run('Аннотация.').italic = True
        p.add_run(f" {rec['abstract'][0]}")

        p = document.add_paragraph('', style='Аннотация')
        p.add_run('Ключевые слова:').italic = True
        p.add_run(f" {rec['kw'][0]}")

        fio = get_fio(rec, lang='ru')
        cit = self.get_cit(rec, fio, lang='ru')
        p = document.add_paragraph('', style='Аннотация')
        p.add_run('Ссылка для цитирования:').italic = True
        p.add_run(f" {cit}")

        document.add_paragraph(f"Текст статьи", style='art_text')
        document.add_paragraph(f"Источники", style='art_section')
        document.add_paragraph(f"Список\nисточников", style='references')

        if authors_num == 1:
            document.add_paragraph(f"Сведения об авторе", style='Аннотация').bold = True
        else:
            document.add_paragraph(f"Сведения об авторах", style='Аннотация').bold = True
        # about_author_n_art = get_about_author_n_art(rec, affs, lang='ru')
        about_author_n_art = get_about_author_n_art2(rec, lang='ru')
        document.add_paragraph(f"{about_author_n_art}", style='about_author')

        # en
        # authors, authors_num = get_full_name(rec, lang='en')
        # affs = get_affs(rec, lang='en')
        about = get_about(rec, lang='en')

        authors, authors_num, affs = authors_affiliation(rec, lang='en')

        document.add_paragraph(f"{authors}", style='art_author')
        document.add_paragraph(f"{affs}\n{about}", style='art_aff')
        document.add_paragraph(rec['title'][1], style='art_title')

        p = document.add_paragraph('', style='Аннотация')
        p.add_run('Abstract.').italic = True
        p.add_run(f" {rec['abstract'][1]}")

        p = document.add_paragraph('', style='Аннотация')
        p.add_run('Keywords:').italic = True
        p.add_run(f" {rec['kw'][1]}")

        fio = get_fio(rec, lang='en')
        cit = self.get_cit(rec, fio, lang='en')
        p = document.add_paragraph('', style='Аннотация')
        p.add_run('For citation:').italic = True
        p.add_run(f" {cit}")

        document.add_paragraph(f"References", style='art_section')
        document.add_paragraph(f"Список\nисточников", style='references')

        if authors_num == 1:
            document.add_paragraph(f"Information about the author", style='Аннотация').bold = True
        else:
            document.add_paragraph(f"Information about the authors", style='Аннотация').bold = True
        # about_author_n_art = get_about_author_n_art(rec, affs, lang='en')
        about_author_n_art = get_about_author_n_art2(rec, lang='en')
        document.add_paragraph(f"{about_author_n_art}", style='about_author')

        document.add_page_break()

    def make_docx(self, issue):
        doc = Document('data/pst_tmpl.docx')
        # doc.add_heading('СОДЕРЖАНИЕ', level=2)
        issue_edn = self.edn[self.issue][0]
        doc.add_paragraph(f'EDN: {issue_edn}', style='Аннотация')
        doc.add_page_break()

        doc.add_paragraph('СОДЕРЖАНИЕ', style='section')
        art_in_iss = self.art_in_iss
        ## содержание
        content_ru = []
        for i in art_in_iss:
            fio = get_fio(i, lang='ru')
            content_ru.append((i['section']['sect_name'][0], fio, i['title'][0]))
        add_content(doc, content_ru)

        # doc.add_heading('CONTENTS', level=2)
        doc.add_paragraph('CONTENTS', style='section')
        content_en = []
        for i in art_in_iss:
            fio = get_fio(i, lang='en')
            content_en.append((i['section']['sect_name'][1], fio, i['title'][1]))
        add_content(doc, content_en)

        ## добавление статей
        section = None
        for a in art_in_iss:
            curr_section = a['section']['sect_name'][0]
            if section == curr_section:
                section = None
            else:
                section = curr_section
            self.add_art(doc, a, section)

        doc.save(f'pst_{issue}.docx')


def get_receiving(rec, lang='ru'):
    receiving = ''
    if lang == 'ru':
        receiving = f"Статья поступила в редакцию: {rec['received']}; "
        if rec['revised']:
            receiving += f"поступила после рецензирования и доработки: {rec['revised']}; "
        receiving += f"принята к публикации: {rec['accepted']}.\n"
    elif lang == 'en':
        receiving = f"Received: {rec['received']}; "
        if rec['revised']:
            receiving += f"revised after review: {rec['revised']}; "
        receiving += f"accepted for publication: {rec['accepted']}.\n"
    return receiving

def get_af(rec, pos):
    af = ''
    for i in rec['authors']:
        for aff in i['affs']:
            af += aff['aff_name'][pos]
            # if aff['position'][pos] is not None:
            if 'position' in aff.keys():
                af += aff['position'][pos]
            else:
                pass
    return af



#
# def get_about_author_n_art(rec, affs, lang='ru'):
#     lst = []
#     if lang == 'ru':
#         pos = 0
#         about_art = get_receiving(rec, lang='ru')
#     elif lang == 'en':
#         pos = 1
#         about_art = get_receiving(rec, lang='en')
#     for i in rec['authors']:
#         # print(i)
#         if i['degree'][pos] is not None:
#             degree = f"{i['degree'][pos]}, "
#         else:
#             degree = ''
#         data = f"{i['surname'][pos]} {i['name'][pos]} {i['patronymic'][pos]}, {degree} \n{affs}, \n{i['about'][pos]}\n"
#         data += f"{i['email']}\n"
#         lst.append(data)
#     lst.append('\n' + about_art)
#     print('get_about_author_n_art', lst)
#     # print(get_af(rec, pos))
#     return '\n'.join(lst)


def get_about_author_n_art2(rec, lang='ru'):
    lst = []
    if lang == 'ru':
        pos = 0
        about_art = get_receiving(rec, lang='ru')
    elif lang == 'en':
        pos = 1
        about_art = get_receiving(rec, lang='en')
    for i in rec['authors']:
        # print(i)
        if i['degree'][pos] is not None:
            degree = f"{i['degree'][pos]},"
        else:
            degree = ''
        af = ''
        for aff in i['affs']:
            if 'position' in aff.keys():
                af += aff['position'][pos]
                af += ', '
            af += aff['aff_name'][pos]
            af += '; '

        af = af.strip('; ')
        data = f"{i['surname'][pos]} {i['name'][pos]} {i['patronymic'][pos]}, {degree} {af}, {i['about'][pos]}. "
        data += f"{i['email']} "
        if 'orcid' in i.keys():
            if i['orcid']:
                data += f"ORCID: {i['orcid']} "
        if 'spin' in i.keys():
            if i['spin']:
                data += f"SPIN: {i['spin']}\n"

        lst.append(data)
    lst.append('\n' + about_art)
    # print('get_about_author_n_art', lst)
    # print(get_af(rec, pos))
    return '\n'.join(lst)


def get_fio(rec, lang='ru'):
    fio_lst = []
    if lang == 'ru':
        pos = 0
        for i in rec['authors']:
            fio = f"{i['surname'][pos]} {i['name'][pos][0]}. {i['patronymic'][pos][0]}."
            fio_lst.append(fio)
    elif lang == 'en':
        pos = 1
        for i in rec['authors']:
            fio = f"{i['surname'][pos]} {i['name'][pos][0]}. {i['patronymic'][pos]}"
            fio_lst.append(fio)
    return ', '.join(fio_lst)


def get_full_name(rec, lang='ru'):
    fio_lst = []
    if lang == 'ru':
        pos = 0
    elif lang == 'en':
        pos = 1
    for i in rec['authors']:
        fio = f"{i['name'][pos]} {i['patronymic'][pos]} {i['surname'][pos]}"
        fio_lst.append(fio)
    return ', '.join(fio_lst), len(fio_lst)

def get_affs(rec, lang='ru'):
    aff_lst = []
    if lang == 'ru':
        pos = 0
    elif lang == 'en':
        pos = 1
    for i in rec['authors']:
        for a in i['affs']:
            aff = f"{a['aff_name'][pos]}"
            aff_lst.append(aff)
    return ', '.join(set(aff_lst))


def get_about(rec, lang='ru'):
    lst = []
    if lang == 'ru':
        pos = 0
    elif lang == 'en':
        pos = 1
    for i in rec['authors']:
        about = f"{i['about'][pos]}"
        lst.append(about)
    return ', '.join(set(lst))


def add_content(document, content):
    section = ''
    for c in content:
        if c[0] != section:
            section = c[0]
            # document.add_paragraph(section)
            document.add_paragraph(section, style='cont_section')
        # p = document.add_paragraph('')
        p = document.add_paragraph('', style='cont_item')
        p.add_run(c[1]).italic = True
        p.add_run(f" {c[2]}")
    document.add_page_break()




def authors_affiliation(rec, lang='ru'):
    print('authors_affiliation')
    lst = []
    if lang == 'ru':
        pos = 0
    elif lang == 'en':
        pos = 1
    authors, authors_num = get_full_name(rec, lang=lang)
    print(authors, authors_num)
    aff_lst = []
    for i in rec['authors']:
        for a in i['affs']:
            aff = f"{a['aff_name'][pos]}"
            aff_lst.append(aff)
    aff_lst = [[a['aff_name'][pos] for a in i['affs']] for i in rec['authors']]
    print(aff_lst)
    # # return ', '.join(set(aff_lst))
    uniq_aff = set([j for i in aff_lst for j in i])
    uniq_aff_seq = []
    for i in aff_lst:
        for j in i:
            if j not in uniq_aff_seq:
                uniq_aff_seq.append(j)
    print('uniq_aff_seq', uniq_aff_seq)
    if authors_num == 1:
        return authors, authors_num, ', '.join([j for i in aff_lst for j in i])
    else:
        dct = {i : [] for i in uniq_aff}
        print(dct)
        authors_w_num = []
        for n, author in enumerate(authors.split(', ')):
            authors_w_num.append(f"{author}{n+1}")
            for affil in aff_lst[n]:
                dct[affil].append(n+1)
        print(dct)

        aff_w_num = []
        for n, af in enumerate(uniq_aff_seq):
            a_aff = []
            for m, au in enumerate(authors.split(', ')):
                print(m)
                a_af_lst = aff_lst[m]
                for a_af in a_af_lst:
                    if a_af == af:
                        print(n, af, m, au, a_af)
                        a_aff.append(m+1)
            aff_w_num.append([af, a_aff])

        au_w_num = []

        for m, au in enumerate(authors.split(', ')):
            au_aff_lst = []
            for n, af in enumerate(uniq_aff_seq):
                a_af_lst = aff_lst[m]
                for a_af in a_af_lst:
                    if a_af == af:
                        # print(n, af, m, au, a_af)
                        au_aff_lst.append(n+1)
            au_w_num.append([au, au_aff_lst])


        print('aff_w_num', aff_w_num)
        print([','.join([str(num) for num in i[1]]) + i[0] for i in aff_w_num])
        aff_w_num_str = [','.join([str(num) for num in i[1]]) + i[0] for i in aff_w_num]
        au_w_num_str = [i[0] + ','.join([str(num) for num in i[1]]) for i in au_w_num]

        affil_w_num = ',\n'.join([str(n+1) + i for n, i in enumerate(uniq_aff_seq)])
        print('affil_w_num', affil_w_num)

        return ', '.join(au_w_num_str), authors_num, affil_w_num




